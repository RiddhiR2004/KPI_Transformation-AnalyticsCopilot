from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas

from app.models import FunctionalSpecItem, FunctionalSpecification, BusinessContext


def set_cell_shading(cell: Any, color_hex: str) -> None:
    """Applies background color to a docx table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))


def set_cell_margins(cell: Any, top: int = 100, bottom: int = 100, left: int = 150, right: int = 150) -> None:
    """Applies internal padding to a docx table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def clean_or_fallback(val: str, fallback: str) -> str:
    cleaned = (val or "").strip()
    lowered = cleaned.lower()
    if not cleaned or lowered in ["not specified", "unknown", "n/a", "none", "tbd", "to be determined", "not provided"]:
        return fallback
    return cleaned


def get_default_reporting(category: str) -> str:
    cat = (category or "").lower()
    if "financial" in cat:
        return "Monthly trended charts, quarterly variance reporting, and margin analysis. Drill down capability by cost center and profit center."
    elif "operational" in cat:
        return "Weekly performance run-charts, daily process monitor scorecards. Comparative tracking against prior 30-day rolling averages."
    elif "strategic" in cat:
        return "C-Suite quarterly scorecards, progress bars against annual target milestones, and executive summaries."
    return "Standard monthly performance dashboard, with trailing 12-month trend charts and period-over-period variance metrics."


def get_default_threshold(category: str) -> str:
    cat = (category or "").lower()
    if "financial" in cat:
        return "Green: Within +/- 2% of budget target. Amber: 2% to 5% variance. Red: > 5% variance or actual spend exceeding budget."
    elif "operational" in cat:
        return "Green: Meets or exceeds 95% operating efficiency. Amber: 90% to 94% efficiency. Red: < 90% operating efficiency."
    elif "strategic" in cat:
        return "Green: Project milestone achieved on time. Amber: 1-2 weeks delay in milestone. Red: > 2 weeks delay in critical path."
    return "Green: Target achieved or exceeded. Amber: 5% to 10% negative variance from target. Red: > 10% negative variance."


def generate_docx_spec(path: Path, spec: Any, context: BusinessContext) -> None:
    """Generates a premium consolidated consulting-grade Word Document for the KPI Functional Specifications."""
    import datetime
    doc = Document()
    
    # Handle list-only fallback
    if isinstance(spec, list):
        items = spec
        spec_status = "DRAFT"
        exec_summary = "Consolidated Executive Summary"
        updated_at_str = datetime.date.today().strftime("%B %d, %Y")
    else:
        items = spec.items
        spec_status = (spec.status or "DRAFT").upper()
        exec_summary = spec.executive_summary or ""
        updated_at_str = spec.updated_at.strftime("%B %d, %Y") if hasattr(spec, "updated_at") and spec.updated_at else datetime.date.today().strftime("%B %d, %Y")

    # Configure page margins (0.75 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Document Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(30, 30, 30)

    # --- Title Page / Cover Section ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("KPI Advisory & Analytics")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(180, 150, 0) # EY Gold

    main_title_p = doc.add_paragraph()
    main_title_p.paragraph_format.space_after = Pt(24)
    main_title_run = main_title_p.add_run("Functional Specification Document")
    main_title_run.font.name = 'Calibri'
    main_title_run.font.size = Pt(28)
    main_title_run.font.bold = True
    main_title_run.font.color.rgb = RGBColor(27, 27, 27)

    # Accent yellow bar under title
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.autofit = False
    accent_table.columns[0].width = Inches(7.0)
    accent_cell = accent_table.cell(0, 0)
    set_cell_shading(accent_cell, "FFE600") # EY Yellow
    set_cell_margins(accent_cell, top=30, bottom=30, left=0, right=0)
    
    # Subtitle / description
    desc_p = doc.add_paragraph()
    desc_p.paragraph_format.space_before = Pt(36)
    desc_p.paragraph_format.space_after = Pt(120)
    desc_run = desc_p.add_run("A unified blueprint translating business strategy into governed, measurable performance metrics.")
    desc_run.font.name = 'Calibri'
    desc_run.font.size = Pt(12)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # --- Document Metadata Section ---
    meta_title = doc.add_paragraph()
    meta_title.paragraph_format.space_before = Pt(12)
    meta_title.paragraph_format.space_after = Pt(12)
    meta_title_run = meta_title.add_run("Document Control & Metadata")
    meta_title_run.font.size = Pt(16)
    meta_title_run.font.bold = True
    meta_title_run.font.color.rgb = RGBColor(27, 27, 27)

    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = 'Table Grid'
    meta_fields = [
        ("Document Version", "1.0"),
        ("Generated Date", updated_at_str),
        ("Industry", clean_or_fallback(context.industry, "Not Specified")),
        ("Organizational Level", clean_or_fallback(context.organization_level, "Not Specified")),
        ("Number of KPIs", f"{len(items)} Approved Performance Metrics"),
        ("Approval Status", spec_status)
    ]
    for i, (label, val) in enumerate(meta_fields):
        row = meta_table.rows[i]
        
        cell_lbl = row.cells[0]
        cell_lbl.text = label
        set_cell_shading(cell_lbl, "F0F0F0")
        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        cell_lbl.paragraphs[0].runs[0].font.bold = True
        cell_lbl.paragraphs[0].runs[0].font.size = Pt(9.5)
        cell_lbl.width = Inches(2.2)

        cell_val = row.cells[1]
        cell_val.text = str(val)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)
        cell_val.paragraphs[0].runs[0].font.size = Pt(9.5)
        cell_val.width = Inches(4.8)

    doc.add_page_break()

    # --- Section 1: Executive Summary ---
    exec_title = doc.add_paragraph()
    exec_title.paragraph_format.space_before = Pt(12)
    exec_title.paragraph_format.space_after = Pt(12)
    exec_title_run = exec_title.add_run("1. Executive Summary")
    exec_title_run.font.size = Pt(16)
    exec_title_run.font.bold = True
    exec_title_run.font.color.rgb = RGBColor(27, 27, 27)

    exec_p = doc.add_paragraph()
    exec_p.paragraph_format.space_after = Pt(18)
    exec_p.paragraph_format.line_spacing = 1.15
    exec_p.add_run(exec_summary or "This document outlines the functional specifications for the approved performance metrics of the organization. Each metric is mapped to strategic objectives and key result areas to ensure operational alignment and governance.")

    # --- Section 2: KPI Landscape Overview ---
    land_title = doc.add_paragraph()
    land_title.paragraph_format.space_before = Pt(18)
    land_title.paragraph_format.space_after = Pt(12)
    land_title.paragraph_format.keep_with_next = True
    land_run = land_title.add_run("2. KPI Landscape Overview")
    land_run.font.size = Pt(16)
    land_run.font.bold = True
    land_run.font.color.rgb = RGBColor(27, 27, 27)

    land_intro = doc.add_paragraph()
    land_intro.paragraph_format.space_after = Pt(12)
    land_intro.add_run("The following visual mind-map displays the KPI Landscape as a branching tree structure, routing from the core KPI Library through strategic categories down to individual approved metrics.")

    # Branching Tree Structure
    p_root = doc.add_paragraph()
    p_root.paragraph_format.left_indent = Inches(0.2)
    p_root.paragraph_format.space_after = Pt(2)
    r_root = p_root.add_run("▼ KPI LIBRARY")
    r_root.font.bold = True
    r_root.font.color.rgb = RGBColor(180, 150, 0) # EY Gold

    categories = {}
    for item in items:
        cat = item.kpi_category or "Operational"
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(item)

    for cat_name, cat_items in categories.items():
        p_cat = doc.add_paragraph()
        p_cat.paragraph_format.left_indent = Inches(0.5)
        p_cat.paragraph_format.space_after = Pt(2)
        r_cat = p_cat.add_run(f" └── 📁 {cat_name.upper()}")
        r_cat.font.bold = True
        r_cat.font.color.rgb = RGBColor(27, 27, 27)

        for item in cat_items:
            p_kpi = doc.add_paragraph()
            p_kpi.paragraph_format.left_indent = Inches(0.9)
            p_kpi.paragraph_format.space_after = Pt(2)
            p_kpi.add_run(f"      ├── 📊 {item.kpi_name}")

    spacer_p = doc.add_paragraph()
    spacer_p.paragraph_format.space_after = Pt(12)

    land_table_intro = doc.add_paragraph()
    land_table_intro.paragraph_format.space_after = Pt(12)
    land_table_intro.add_run("The following table provides a high-level catalog of all approved performance indicators within the scope of this transformation initiative.")

    land_table = doc.add_table(rows=1 + len(items), cols=4)
    land_table.style = 'Table Grid'
    
    headers = ["KPI ID", "KPI Name", "Category", "Functional Area"]
    for col_idx, text in enumerate(headers):
        cell = land_table.cell(0, col_idx)
        cell.text = text
        set_cell_shading(cell, "1B1B1B")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    for row_idx, item in enumerate(items, start=1):
        row = land_table.rows[row_idx]
        
        c0 = row.cells[0]
        c0.text = f"KPI-{row_idx:03d}"
        set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
        c0.paragraphs[0].runs[0].font.size = Pt(9.5)
        
        c1 = row.cells[1]
        c1.text = item.kpi_name
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        c1.paragraphs[0].runs[0].font.size = Pt(9.5)
        c1.paragraphs[0].runs[0].font.bold = True
        
        c2 = row.cells[2]
        c2.text = clean_or_fallback(item.kpi_category, "Operational")
        set_cell_margins(c2, top=60, bottom=60, left=100, right=100)
        c2.paragraphs[0].runs[0].font.size = Pt(9.5)
        
        c3 = row.cells[3]
        c3.text = clean_or_fallback(item.functional_area, "Operations")
        set_cell_margins(c3, top=60, bottom=60, left=100, right=100)
        c3.paragraphs[0].runs[0].font.size = Pt(9.5)

    # --- Section 3: Strategic Traceability Matrix ---
    trace_title = doc.add_paragraph()
    trace_title.paragraph_format.space_before = Pt(18)
    trace_title.paragraph_format.space_after = Pt(12)
    trace_title.paragraph_format.keep_with_next = True
    trace_run = trace_title.add_run("3. Strategic Traceability Matrix")
    trace_run.font.size = Pt(16)
    trace_run.font.bold = True
    trace_run.font.color.rgb = RGBColor(27, 27, 27)

    trace_intro = doc.add_paragraph()
    trace_intro.paragraph_format.space_after = Pt(12)
    trace_intro.add_run("This matrix illustrates the strategic alignment from executive objectives down to specific key performance indicators, providing visibility into strategic translation.")

    trace_table = doc.add_table(rows=1 + len(items), cols=5)
    trace_table.style = 'Table Grid'
    
    trace_headers = ["Strategic Objective", "Business Challenge", "KRA", "Functional Area", "KPI Name"]
    for col_idx, text in enumerate(trace_headers):
        cell = trace_table.cell(0, col_idx)
        cell.text = text
        set_cell_shading(cell, "1B1B1B")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    for row_idx, item in enumerate(items, start=1):
        row = trace_table.rows[row_idx]
        
        c0 = row.cells[0]
        c0.text = clean_or_fallback(item.strategic_objective_supported, "Optimize Strategy")
        set_cell_margins(c0, top=60, bottom=60, left=80, right=80)
        c0.paragraphs[0].runs[0].font.size = Pt(8.5)
        
        c1 = row.cells[1]
        c1.text = clean_or_fallback(item.business_challenge_addressed, "Inefficient Processes")
        set_cell_margins(c1, top=60, bottom=60, left=80, right=80)
        c1.paragraphs[0].runs[0].font.size = Pt(8.5)
        
        c2 = row.cells[2]
        c2.text = clean_or_fallback(item.related_kra, "Operational Excellence")
        set_cell_margins(c2, top=60, bottom=60, left=80, right=80)
        c2.paragraphs[0].runs[0].font.size = Pt(8.5)
        
        c3 = row.cells[3]
        c3.text = clean_or_fallback(item.functional_area, "Operations")
        set_cell_margins(c3, top=60, bottom=60, left=80, right=80)
        c3.paragraphs[0].runs[0].font.size = Pt(8.5)
        
        c4 = row.cells[4]
        c4.text = item.kpi_name
        set_cell_margins(c4, top=60, bottom=60, left=80, right=80)
        c4.paragraphs[0].runs[0].font.bold = True
        c4.paragraphs[0].runs[0].font.size = Pt(8.5)

    # --- Section 4: Individual KPI Sections ---
    doc.add_page_break()
    kpi_sec_title = doc.add_paragraph()
    kpi_sec_title.paragraph_format.space_before = Pt(12)
    kpi_sec_title.paragraph_format.space_after = Pt(12)
    kpi_sec_title_run = kpi_sec_title.add_run("4. Individual KPI Specifications")
    kpi_sec_title_run.font.size = Pt(16)
    kpi_sec_title_run.font.bold = True
    kpi_sec_title_run.font.color.rgb = RGBColor(27, 27, 27)

    for index, item in enumerate(items, start=1):
        kpi_p = doc.add_paragraph()
        kpi_p.paragraph_format.space_before = Pt(24)
        kpi_p.paragraph_format.space_after = Pt(8)
        kpi_p.paragraph_format.keep_with_next = True
        
        num_run = kpi_p.add_run(f"Metric {index:02d}: ")
        num_run.font.size = Pt(13)
        num_run.font.bold = True
        num_run.font.color.rgb = RGBColor(180, 150, 0)
        
        name_run = kpi_p.add_run(item.kpi_name)
        name_run.font.size = Pt(13)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(27, 27, 27)

        # Quick Reference and Owner Block
        kpi_meta_tbl = doc.add_table(rows=2, cols=4)
        kpi_meta_tbl.style = 'Table Grid'
        
        kpi_meta_fields = [
            ("KPI Category", clean_or_fallback(item.kpi_category, "Operational")),
            ("Functional Area", clean_or_fallback(item.functional_area, "Operations")),
            ("Business Owner", clean_or_fallback(item.business_owner, "Advisory Business Owner")),
            ("Data Owner", clean_or_fallback(item.data_owner, "Advisory Data Owner"))
        ]
        
        row_lbls = kpi_meta_tbl.rows[0]
        row_vals = kpi_meta_tbl.rows[1]
        
        for col_idx, (lbl, val) in enumerate(kpi_meta_fields):
            c_lbl = row_lbls.cells[col_idx]
            c_lbl.text = lbl
            set_cell_shading(c_lbl, "F0F0F0")
            set_cell_margins(c_lbl, top=40, bottom=40, left=60, right=60)
            c_lbl.paragraphs[0].runs[0].font.bold = True
            c_lbl.paragraphs[0].runs[0].font.size = Pt(8.5)
            
            c_val = row_vals.cells[col_idx]
            c_val.text = val
            set_cell_margins(c_val, top=40, bottom=40, left=60, right=60)
            c_val.paragraphs[0].runs[0].font.size = Pt(8.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Core details
        details = [
            ("Business Definition", clean_or_fallback(item.kpi_definition, "Official definition of the KPI.")),
            ("Business Purpose & Strategic Relevance", clean_or_fallback(item.business_purpose_relevance, "Detailing the strategic necessity of tracking this metric.")),
        ]
        for title, val in details:
            sh_p = doc.add_paragraph()
            sh_p.paragraph_format.space_before = Pt(6)
            sh_p.paragraph_format.space_after = Pt(2)
            sh_p.paragraph_format.keep_with_next = True
            sh_run = sh_p.add_run(title)
            sh_run.font.bold = True
            sh_run.font.size = Pt(10)
            sh_run.font.color.rgb = RGBColor(80, 80, 80)
            
            p_text = doc.add_paragraph()
            p_text.paragraph_format.space_after = Pt(4)
            p_text.paragraph_format.line_spacing = 1.15
            p_text.add_run(val)
            p_text.runs[0].font.size = Pt(9.5)

        # Calculation Methodology
        calc_title_p = doc.add_paragraph()
        calc_title_p.paragraph_format.space_before = Pt(6)
        calc_title_p.paragraph_format.space_after = Pt(4)
        calc_title_p.paragraph_format.keep_with_next = True
        calc_title_run = calc_title_p.add_run("Calculation Methodology")
        calc_title_run.font.bold = True
        calc_title_run.font.size = Pt(10)
        calc_title_run.font.color.rgb = RGBColor(80, 80, 80)

        calc_fields = [
            ("Formula", clean_or_fallback(item.formula, "Formula not defined.")),
            ("Numerator Details", clean_or_fallback(item.numerator, "Numerator details.")),
            ("Denominator Details", clean_or_fallback(item.denominator, "Denominator details.")),
            ("Calculation Logic", clean_or_fallback(item.calculation_methodology, "Standard calculation procedures.")),
            ("Inclusion Rules", clean_or_fallback(item.inclusion_rules, "Include all standard transactions.")),
            ("Exclusion Rules", clean_or_fallback(item.exclusion_rules, "Exclude test and non-operational entries.")),
            ("Sample Worked Example", clean_or_fallback(item.sample_calculation, "Sample: (Subset / Total) * 100"))
        ]
        
        calc_tbl = doc.add_table(rows=len(calc_fields), cols=2)
        calc_tbl.style = 'Table Grid'
        for r_idx, (lbl, val) in enumerate(calc_fields):
            row = calc_tbl.rows[r_idx]
            
            c_lbl = row.cells[0]
            c_lbl.text = lbl
            set_cell_shading(c_lbl, "F9F9F9")
            set_cell_margins(c_lbl, top=40, bottom=40, left=80, right=80)
            c_lbl.paragraphs[0].runs[0].font.bold = True
            c_lbl.paragraphs[0].runs[0].font.size = Pt(9)
            c_lbl.width = Inches(2.0)
            
            c_val = row.cells[1]
            c_val.text = val
            set_cell_margins(c_val, top=40, bottom=40, left=80, right=80)
            c_val.paragraphs[0].runs[0].font.size = Pt(9)
            c_val.width = Inches(5.0)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Business Rules & Validation
        rules_title_p = doc.add_paragraph()
        rules_title_p.paragraph_format.space_before = Pt(6)
        rules_title_p.paragraph_format.space_after = Pt(4)
        rules_title_p.paragraph_format.keep_with_next = True
        rules_title_run = rules_title_p.add_run("Business Rules & Data Validation")
        rules_title_run.font.bold = True
        rules_title_run.font.size = Pt(10)
        rules_title_run.font.color.rgb = RGBColor(80, 80, 80)

        rules_fields = [
            ("Business Operating Rules", clean_or_fallback(item.business_rules, "Standard business operating rules.")),
            ("Data Validation Rules", clean_or_fallback(item.data_validation_rules, "Ensure numeric bounds and data integrity.")),
            ("Exception Handling Rules", clean_or_fallback(item.exception_handling_rules, "Standard null/zero denominator exceptions.")),
            ("Data Quality Expectations", clean_or_fallback(item.data_quality_expectations, "High accuracy in source ledger."))
        ]
        
        rules_tbl = doc.add_table(rows=len(rules_fields), cols=2)
        rules_tbl.style = 'Table Grid'
        for r_idx, (lbl, val) in enumerate(rules_fields):
            row = rules_tbl.rows[r_idx]
            
            c_lbl = row.cells[0]
            c_lbl.text = lbl
            set_cell_shading(c_lbl, "F9F9F9")
            set_cell_margins(c_lbl, top=40, bottom=40, left=80, right=80)
            c_lbl.paragraphs[0].runs[0].font.bold = True
            c_lbl.paragraphs[0].runs[0].font.size = Pt(9)
            c_lbl.width = Inches(2.0)
            
            c_val = row.cells[1]
            c_val.text = val
            set_cell_margins(c_val, top=40, bottom=40, left=80, right=80)
            c_val.paragraphs[0].runs[0].font.size = Pt(9)
            c_val.width = Inches(5.0)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Recommended Source Systems
        source_title_p = doc.add_paragraph()
        source_title_p.paragraph_format.space_before = Pt(6)
        source_title_p.paragraph_format.space_after = Pt(2)
        source_title_p.paragraph_format.keep_with_next = True
        source_title_run = source_title_p.add_run("Recommended Source Systems")
        source_title_run.font.bold = True
        source_title_run.font.size = Pt(10)
        source_title_run.font.color.rgb = RGBColor(80, 80, 80)

        source_p = doc.add_paragraph()
        source_p.paragraph_format.space_after = Pt(12)
        source_run = source_p.add_run(clean_or_fallback(item.source_systems_lineage, "Recommended source ERP database. (Assumption)"))
        source_run.font.size = Pt(9.5)
        source_run.font.italic = True

    # --- Section 5: Governance Framework ---
    doc.add_page_break()
    gov_title = doc.add_paragraph()
    gov_title.paragraph_format.space_before = Pt(12)
    gov_title.paragraph_format.space_after = Pt(12)
    gov_title_run = gov_title.add_run("5. Governance Framework")
    gov_title_run.font.size = Pt(16)
    gov_title_run.font.bold = True
    gov_title_run.font.color.rgb = RGBColor(27, 27, 27)

    gov_intro = doc.add_paragraph()
    gov_intro.paragraph_format.space_after = Pt(12)
    gov_intro.add_run(
        "To ensure metric consistency, accountability, and ongoing relevance, a formal governance framework is "
        "established for all approved KPIs. This structure assigns clear responsibilities and defines escalation paths."
    )

    # Roles and Responsibilities sub-section
    roles_title = doc.add_paragraph()
    roles_title.paragraph_format.space_before = Pt(8)
    roles_title.paragraph_format.space_after = Pt(6)
    roles_title_run = roles_title.add_run("Roles and Responsibilities")
    roles_title_run.font.bold = True
    roles_title_run.font.size = Pt(12)
    roles_title_run.font.color.rgb = RGBColor(80, 80, 80)

    roles_text = [
        ("Business Owner", "Responsible for defining the business logic, validating calculation results, approving target thresholds, and driving operational performance based on metric insights."),
        ("Data Owner", "Accountable for technical lineage, data completeness, source-to-target mapping, ETL data quality checks, and resolving data ingestion or availability issues."),
        ("Escalation Path", "In case of data quality discrepancies or alignment disputes, issues are escalated to the Data Governance Committee and KPI Advisory Board for review and reconciliation.")
    ]
    for role, desc in roles_text:
        role_p = doc.add_paragraph(style='List Bullet')
        role_p.paragraph_format.space_after = Pt(4)
        r_run = role_p.add_run(f"{role}: ")
        r_run.bold = True
        role_p.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Ownership & Governance Matrix
    matrix_title = doc.add_paragraph()
    matrix_title.paragraph_format.space_before = Pt(8)
    matrix_title.paragraph_format.space_after = Pt(6)
    matrix_title_run = matrix_title.add_run("Ownership and Governance Policy Matrix")
    matrix_title_run.font.bold = True
    matrix_title_run.font.size = Pt(12)
    matrix_title_run.font.color.rgb = RGBColor(80, 80, 80)

    gov_tbl = doc.add_table(rows=1 + len(items), cols=4)
    gov_tbl.style = 'Table Grid'
    
    gov_headers = ["Metric Name", "Business Owner", "Data Owner", "Governance Policy Notes"]
    for col_idx, text in enumerate(gov_headers):
        cell = gov_tbl.cell(0, col_idx)
        cell.text = text
        set_cell_shading(cell, "1B1B1B")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, item in enumerate(items, start=1):
        row = gov_tbl.rows[row_idx]
        
        c0 = row.cells[0]
        c0.text = item.kpi_name
        set_cell_margins(c0, top=60, bottom=60, left=80, right=80)
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(9)
        
        c1 = row.cells[1]
        c1.text = clean_or_fallback(item.business_owner, "Business Sponsor")
        set_cell_margins(c1, top=60, bottom=60, left=80, right=80)
        c1.paragraphs[0].runs[0].font.size = Pt(9)
        
        c2 = row.cells[2]
        c2.text = clean_or_fallback(item.data_owner, "Data Custodian")
        set_cell_margins(c2, top=60, bottom=60, left=80, right=80)
        c2.paragraphs[0].runs[0].font.size = Pt(9)
        
        c3 = row.cells[3]
        c3.text = clean_or_fallback(item.ownership_governance, "Subject to standard quarterly advisory audits and performance reviews.")
        set_cell_margins(c3, top=60, bottom=60, left=80, right=80)
        c3.paragraphs[0].runs[0].font.size = Pt(8.5)

    # --- Section 6: Reporting & Dashboard Requirements ---
    doc.add_page_break()
    rep_title = doc.add_paragraph()
    rep_title.paragraph_format.space_before = Pt(12)
    rep_title.paragraph_format.space_after = Pt(12)
    rep_title_run = rep_title.add_run("6. Reporting & Dashboard Requirements")
    rep_title_run.font.size = Pt(16)
    rep_title_run.font.bold = True
    rep_title_run.font.color.rgb = RGBColor(27, 27, 27)

    rep_intro = doc.add_paragraph()
    rep_intro.paragraph_format.space_after = Pt(12)
    rep_intro.add_run(
        "Visualization layout and threshold criteria dictate how data is displayed to support decision-making. "
        "The matrix below outlines reporting recommendations and performance thresholds for each metric."
    )

    rep_tbl = doc.add_table(rows=1 + len(items), cols=4)
    rep_tbl.style = 'Table Grid'
    
    rep_headers = ["Metric Name", "Reporting Guidelines", "Dashboard Placement", "Threshold Guidance"]
    for col_idx, text in enumerate(rep_headers):
        cell = rep_tbl.cell(0, col_idx)
        cell.text = text
        set_cell_shading(cell, "1B1B1B")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, item in enumerate(items, start=1):
        row = rep_tbl.rows[row_idx]
        
        def_rep = get_default_reporting(item.kpi_category)
        def_thresh = get_default_threshold(item.kpi_category)
        def_dash = "Standard Performance Dashboard."
        
        rep_guidelines = clean_or_fallback(item.reporting_requirements, def_rep)
        dash_placement = clean_or_fallback(item.dashboard_recommendations, def_dash)
        threshold_guidance = clean_or_fallback(item.threshold_guidance, def_thresh)
        
        c0 = row.cells[0]
        c0.text = item.kpi_name
        set_cell_margins(c0, top=60, bottom=60, left=80, right=80)
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(9)
        
        c1 = row.cells[1]
        c1.text = rep_guidelines
        set_cell_margins(c1, top=60, bottom=60, left=80, right=80)
        c1.paragraphs[0].runs[0].font.size = Pt(8.5)
        
        c2 = row.cells[2]
        c2.text = dash_placement
        set_cell_margins(c2, top=60, bottom=60, left=80, right=80)
        c2.paragraphs[0].runs[0].font.size = Pt(8.5)
        
        c3 = row.cells[3]
        c3.text = threshold_guidance
        set_cell_margins(c3, top=60, bottom=60, left=80, right=80)
        c3.paragraphs[0].runs[0].font.size = Pt(8.5)

    # --- Section 7: Assumptions & Constraints ---
    doc.add_page_break()
    asmp_title = doc.add_paragraph()
    asmp_title.paragraph_format.space_before = Pt(12)
    asmp_title.paragraph_format.space_after = Pt(12)
    asmp_title_run = asmp_title.add_run("7. Assumptions & Constraints")
    asmp_title_run.font.size = Pt(16)
    asmp_title_run.font.bold = True
    asmp_title_run.font.color.rgb = RGBColor(27, 27, 27)

    asmp_intro = doc.add_paragraph()
    asmp_intro.paragraph_format.space_after = Pt(12)
    asmp_intro.add_run(
        "A clear understanding of business and technical assumptions is critical for successful implementation. "
        "The following list represents consolidated baseline assumptions and constraints."
    )

    gen_title = doc.add_paragraph()
    gen_title.paragraph_format.space_before = Pt(8)
    gen_title.paragraph_format.space_after = Pt(6)
    gen_title_run = gen_title.add_run("General Architectural Assumptions")
    gen_title_run.font.bold = True
    gen_title_run.font.size = Pt(12)
    gen_title_run.font.color.rgb = RGBColor(80, 80, 80)

    gen_asmps = [
        "Data Availability: Source transactional tables are assumed to be loaded into the central reporting repository on a standard nightly batch cadence.",
        "Fiscal Calendar: Standard calendar year rules are assumed unless otherwise explicitly documented in specific financial indicators.",
        "System Uptime: Target source ERP ledgers are assumed to maintain 99.5% uptime during standard reporting extraction windows."
    ]
    for asm in gen_asmps:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        lbl, desc = asm.split(":", 1)
        r = p.add_run(lbl + ":")
        r.bold = True
        p.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    spec_title = doc.add_paragraph()
    spec_title.paragraph_format.space_before = Pt(8)
    spec_title.paragraph_format.space_after = Pt(6)
    spec_title_run = spec_title.add_run("Metric-Specific Assumptions and Limitations")
    spec_title_run.font.bold = True
    spec_title_run.font.size = Pt(12)
    spec_title_run.font.color.rgb = RGBColor(80, 80, 80)

    for item in items:
        cleaned_asm = clean_or_fallback(item.assumptions_constraints, "").strip()
        if cleaned_asm:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(f"{item.kpi_name}: ")
            r.bold = True
            p.add_run(cleaned_asm)

    # --- Section 8: Implementation Considerations ---
    doc.add_page_break()
    impl_title = doc.add_paragraph()
    impl_title.paragraph_format.space_before = Pt(12)
    impl_title.paragraph_format.space_after = Pt(12)
    impl_title_run = impl_title.add_run("8. Implementation Considerations")
    impl_title_run.font.size = Pt(16)
    impl_title_run.font.bold = True
    impl_title_run.font.color.rgb = RGBColor(27, 27, 27)

    impl_intro = doc.add_paragraph()
    impl_intro.paragraph_format.space_after = Pt(12)
    impl_intro.add_run(
        "Transitioning these specifications into functional BI tools requires rigorous testing, change management, "
        "and data reconciliation. Standard implementation considerations are outlined below."
    )

    gen_impl_title = doc.add_paragraph()
    gen_impl_title.paragraph_format.space_before = Pt(8)
    gen_impl_title.paragraph_format.space_after = Pt(6)
    gen_impl_title_run = gen_impl_title.add_run("Standard Implementation Guidelines")
    gen_impl_title_run.font.bold = True
    gen_impl_title_run.font.size = Pt(12)
    gen_impl_title_run.font.color.rgb = RGBColor(80, 80, 80)

    gen_impls = [
        "Data Reconciliation: All KPI calculation outcomes must be audited and reconciled against official books of record or audited financial statements.",
        "User Acceptance Testing (UAT): Business owners must perform visual and numeric validation of dashboard mockups prior to production sign-off.",
        "Change Management: Training sessions and clear system documentation are required to support user onboarding and ensure high organizational adoption."
    ]
    for imp in gen_impls:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        lbl, desc = imp.split(":", 1)
        r = p.add_run(lbl + ":")
        r.bold = True
        p.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    spec_impl_title = doc.add_paragraph()
    spec_impl_title.paragraph_format.space_before = Pt(8)
    spec_impl_title.paragraph_format.space_after = Pt(6)
    spec_impl_title_run = spec_impl_title.add_run("Metric-Specific Technical Considerations")
    spec_impl_title_run.font.bold = True
    spec_impl_title_run.font.size = Pt(12)
    spec_impl_title_run.font.color.rgb = RGBColor(80, 80, 80)

    for item in items:
        cleaned_imp = clean_or_fallback(item.implementation_guidance, "").strip()
        if cleaned_imp:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(f"{item.kpi_name}: ")
            r.bold = True
            p.add_run(cleaned_imp)

    # --- Section 9: Appendix ---
    doc.add_page_break()
    appx_title = doc.add_paragraph()
    appx_title.paragraph_format.space_before = Pt(12)
    appx_title.paragraph_format.space_after = Pt(12)
    appx_title_run = appx_title.add_run("9. Appendix")
    appx_title_run.font.size = Pt(16)
    appx_title_run.font.bold = True
    appx_title_run.font.color.rgb = RGBColor(27, 27, 27)

    glos_title = doc.add_paragraph()
    glos_title.paragraph_format.space_before = Pt(8)
    glos_title.paragraph_format.space_after = Pt(6)
    glos_title_run = glos_title.add_run("KPI Glossary")
    glos_title_run.font.bold = True
    glos_title_run.font.size = Pt(12)
    glos_title_run.font.color.rgb = RGBColor(80, 80, 80)

    glossary_items = [
        ("KPI (Key Performance Indicator)", "A quantifiable measure used to evaluate the success of an organization or activity in meeting performance objectives."),
        ("KRA (Key Result Area)", "Primary focus areas of outcomes or outputs for which an organizational unit or role is responsible."),
        ("Numerator", "The upper portion of a division representing the measured subset of occurrences."),
        ("Denominator", "The lower portion of a division representing the total base population.")
    ]
    for term, defn in glossary_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(term + ": ")
        r.bold = True
        p.add_run(defn)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    dq_title = doc.add_paragraph()
    dq_title.paragraph_format.space_before = Pt(8)
    dq_title.paragraph_format.space_after = Pt(6)
    dq_title_run = dq_title.add_run("Data Quality Principles")
    dq_title_run.font.bold = True
    dq_title_run.font.size = Pt(12)
    dq_title_run.font.color.rgb = RGBColor(80, 80, 80)

    dq_items = [
        ("Accuracy", "Data correctly represents the real-world operational event it records."),
        ("Completeness", "All necessary dataset components are present without omission."),
        ("Consistency", "Metrics align across various systems, business units, and report interfaces."),
        ("Timeliness", "Updates occur within the required reporting cadence and operational windows.")
    ]
    for principle, desc in dq_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(principle + ": ")
        r.bold = True
        p.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    ac_title = doc.add_paragraph()
    ac_title.paragraph_format.space_before = Pt(8)
    ac_title.paragraph_format.space_after = Pt(6)
    ac_title_run = ac_title.add_run("Acronym Reference")
    ac_title_run.font.bold = True
    ac_title_run.font.size = Pt(12)
    ac_title_run.font.color.rgb = RGBColor(80, 80, 80)

    ac_items = [
        ("ERP", "Enterprise Resource Planning"),
        ("SAP FI-CO", "Financial Accounting & Controlling"),
        ("SAP SD", "Sales & Distribution"),
        ("SAP MM", "Materials Management"),
        ("BI / DWH", "Business Intelligence / Data Warehouse"),
        ("UAT", "User Acceptance Testing")
    ]
    for acr, full in ac_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(acr + ": ")
        r.bold = True
        p.add_run(full)

    doc.save(path)




# --- Custom PDF Canvas for ReportLab to support running headers & page numbers ---
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args: Any, **kwargs: Any) -> None:
        super().__init__(*args, **kwargs)
        self._saved_page_states: list[dict[str, Any]] = []

    def showPage(self) -> None:
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self) -> None:
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, total_pages: int) -> None:
        self.saveState()
        
        # Suppress headers/footers on the title page
        if self._pageNumber == 1:
            self.restoreState()
            return

        # Top border accent (EY Yellow)
        self.setFillColor(colors.HexColor("#FFE600"))
        self.rect(0, 782, 612, 10, fill=True, stroke=False)

        # Header Text
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#666666"))
        self.drawRightString(558, 760, "KPI FUNCTIONAL SPECIFICATION DOCUMENT")

        # Top separator line
        self.setStrokeColor(colors.HexColor("#DCDCD9"))
        self.setLineWidth(0.5)
        self.line(54, 750, 558, 750)

        # Bottom separator line & footer text
        self.line(54, 55, 558, 55)
        self.drawString(54, 40, "Confidential - Advisory Work Product")
        self.drawRightString(558, 40, f"Page {self._pageNumber} of {total_pages}")
        
        self.restoreState()


def draw_kpi_tree_pdf(items: list) -> Any:
    from reportlab.graphics.shapes import Drawing, Rect, String, Line, Circle, Group
    from reportlab.lib.colors import HexColor
    
    # Group KPIs by category
    categories = {}
    for item in items:
        cat = item.kpi_category or "Operational"
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(item)
        
    n_kpis = len(items)
    kpi_gap = 30
    height = max(180, n_kpis * kpi_gap + 20)
    width = 504
    
    d = Drawing(width, height)
    
    # Background border for drawing container
    d.add(Rect(0, 0, width, height, fillColor=HexColor("#F9F9F9"), strokeColor=HexColor("#DCDCD9"), strokeWidth=0.5, rx=4, ry=4))
    
    # Coordinates mapping
    root_x = 15
    root_w = 90
    root_h = 30
    
    cat_x = 160
    cat_w = 100
    cat_h = 24
    
    kpi_x = 310
    kpi_w = 180
    kpi_h = 20
    
    # Calculate layouts
    kpis_layout = {}
    cats_layout = {}
    
    current_kpi_idx = 0
    for category, cat_items in categories.items():
        # KPI y positions
        y_positions = []
        for idx, item in enumerate(cat_items):
            y = height - 20 - ((current_kpi_idx + idx) * kpi_gap) - kpi_h/2
            kpis_layout[item.id] = (kpi_x, y)
            y_positions.append(y)
            
        # Category y is the average of its KPIs' y positions
        cat_y = sum(y_positions) / len(y_positions) if y_positions else height / 2
        cats_layout[category] = (cat_x, cat_y)
        current_kpi_idx += len(cat_items)
        
    # Root y is the average of category y positions
    cat_ys = [pos[1] for pos in cats_layout.values()]
    root_y = sum(cat_ys) / len(cat_ys) if cat_ys else height / 2
    
    # Draw connections first so nodes sit on top
    # 1. Root to Categories
    for cat_name, (cx, cy) in cats_layout.items():
        rx1 = root_x + root_w
        ry1 = root_y
        cx1 = cx
        cy1 = cy
        mid_x = (rx1 + cx1) / 2
        d.add(Line(rx1, ry1, mid_x, ry1, strokeColor=HexColor("#FFE600"), strokeWidth=1.5))
        d.add(Line(mid_x, ry1, mid_x, cy1, strokeColor=HexColor("#FFE600"), strokeWidth=1.5))
        d.add(Line(mid_x, cy1, cx1, cy1, strokeColor=HexColor("#FFE600"), strokeWidth=1.5))
        
    # 2. Categories to KPIs
    for item in items:
        cat_name = item.kpi_category or "Operational"
        if cat_name in cats_layout and item.id in kpis_layout:
            cx, cy = cats_layout[cat_name]
            kx, ky = kpis_layout[item.id]
            cx1 = cx + cat_w
            cy1 = cy
            kx1 = kx
            ky1 = ky + kpi_h/2
            mid_x = (cx1 + kx1) / 2
            d.add(Line(cx1, cy1, mid_x, cy1, strokeColor=HexColor("#DCDCD9"), strokeWidth=1.0))
            d.add(Line(mid_x, cy1, mid_x, ky1, strokeColor=HexColor("#DCDCD9"), strokeWidth=1.0))
            d.add(Line(mid_x, ky1, kx1, ky1, strokeColor=HexColor("#DCDCD9"), strokeWidth=1.0))
            
    # Draw Root Node
    d.add(Rect(root_x, root_y - root_h/2, root_w, root_h, fillColor=HexColor("#1B1B1B"), strokeColor=HexColor("#1B1B1B"), rx=4, ry=4))
    d.add(String(root_x + root_w/2, root_y - 3, "KPI LIBRARY", textAnchor="middle", fontName="Helvetica-Bold", fontSize=8, fillColor=HexColor("#FFE600")))
    
    # Draw Category Nodes
    for cat_name, (cx, cy) in cats_layout.items():
        d.add(Rect(cx, cy - cat_h/2, cat_w, cat_h, fillColor=HexColor("#FFE600"), strokeColor=HexColor("#B49600"), strokeWidth=0.5, rx=3, ry=3))
        disp_cat = cat_name
        if len(disp_cat) > 18:
            disp_cat = disp_cat[:15] + "..."
        d.add(String(cx + cat_w/2, cy - 3, disp_cat.upper(), textAnchor="middle", fontName="Helvetica-Bold", fontSize=7, fillColor=HexColor("#1B1B1B")))
        
    # Draw KPI Nodes
    for idx, item in enumerate(items):
        if item.id in kpis_layout:
            kx, ky = kpis_layout[item.id]
            d.add(Rect(kx, ky, kpi_w, kpi_h, fillColor=HexColor("#FFFFFF"), strokeColor=HexColor("#DCDCD9"), strokeWidth=0.5, rx=2, ry=2))
            
            kpi_name = item.kpi_name or ""
            if len(kpi_name) > 30:
                kpi_name = kpi_name[:27] + "..."
            
            d.add(String(kx + 8, ky + 6, f"{idx+1:02d}. {kpi_name}", fontName="Helvetica", fontSize=7.5, fillColor=HexColor("#1B1B1B")))
            
    return d


def generate_pdf_spec(path: Path, spec: Any, context: BusinessContext, doc_name: str | None = None) -> None:
    """Generates a premium consolidated client-ready PDF document utilizing ReportLab's flowable architecture."""
    import datetime
    
    # Handle list-only fallback
    if isinstance(spec, list):
        items = spec
        spec_status = "DRAFT"
        exec_summary = "Consolidated Executive Summary"
        updated_at_str = datetime.date.today().strftime("%B %d, %Y")
    else:
        items = spec.items
        spec_status = (spec.status or "DRAFT").upper()
        exec_summary = spec.executive_summary or ""
        updated_at_str = spec.updated_at.strftime("%B %d, %Y") if hasattr(spec, "updated_at") and spec.updated_at else datetime.date.today().strftime("%B %d, %Y")

    pdf_title = doc_name if doc_name else "KPI Functional Specification Document"
    pdf_author = "KPI Advisory & Analytics"
    pdf_subject = "KPI Functional Specification Document"
    pdf_creator = "KPI Advisory & Analytics Copilot"

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=72,
        bottomMargin=72,
        title=pdf_title,
        author=pdf_author,
        subject=pdf_subject,
        creator=pdf_creator
    )

    styles = getSampleStyleSheet()
    
    # Custom Palette
    ey_yellow = colors.HexColor("#FFE600")
    dark_gray = colors.HexColor("#1B1B1B")
    border_color = colors.HexColor("#DCDCD9")
    text_color = colors.HexColor("#2C2C2A")

    # Typography / Paragraph Styles
    title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=26,
        leading=32,
        textColor=dark_gray,
        spaceAfter=15
    )
    
    subtitle_style = ParagraphStyle(
        'CoverSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#B49600"),
        spaceAfter=10,
        textTransform='uppercase'
    )

    body_style = ParagraphStyle(
        'SpecBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=text_color,
        spaceAfter=6
    )

    section_heading = ParagraphStyle(
        'SectionHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=15,
        leading=19,
        textColor=dark_gray,
        spaceBefore=16,
        spaceAfter=10,
        keepWithNext=True
    )

    kpi_title_style = ParagraphStyle(
        'KpiTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=dark_gray,
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True
    )

    sub_section_title_style = ParagraphStyle(
        'SubSectionTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#555555"),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )

    tbl_label_style = ParagraphStyle(
        'TblLabel',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11,
        textColor=dark_gray
    )

    tbl_value_style = ParagraphStyle(
        'TblValue',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=11,
        textColor=text_color
    )

    tbl_hdr_style = ParagraphStyle(
        'TblHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11,
        textColor=colors.white
    )

    bullet_style = ParagraphStyle(
        'BulletPoint',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        leftIndent=20,
        firstLineIndent=-10,
        textColor=text_color,
        spaceAfter=4
    )

    story = []

    # --- Cover Page Layout ---
    story.append(Spacer(1, 100))
    story.append(Paragraph("KPI Advisory & Analytics", subtitle_style))
    story.append(Paragraph("Functional Specification Document", title_style))
    
    # Yellow colored accent bar
    bar_data = [['']]
    bar_table = Table(bar_data, colWidths=[504], rowHeights=[4])
    bar_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), ey_yellow),
        ('PADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ('TOPPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(bar_table)
    story.append(Spacer(1, 24))
    story.append(Paragraph("A unified blueprint translating business strategy into governed, measurable performance metrics.", body_style))
    story.append(PageBreak())

    # --- Document Control & Metadata ---
    story.append(Paragraph("Document Control & Metadata", section_heading))
    story.append(Spacer(1, 10))
    
    meta_data = [
        [Paragraph("Document Version", tbl_label_style), Paragraph("1.0", tbl_value_style)],
        [Paragraph("Generated Date", tbl_label_style), Paragraph(updated_at_str, tbl_value_style)],
        [Paragraph("Industry Sector", tbl_label_style), Paragraph(clean_or_fallback(context.industry, "Not Specified"), tbl_value_style)],
        [Paragraph("Organizational Level", tbl_label_style), Paragraph(clean_or_fallback(context.organization_level, "Not Specified"), tbl_value_style)],
        [Paragraph("Scope of Performance Metrics", tbl_label_style), Paragraph(f"{len(items)} Approved Performance Metrics", tbl_value_style)],
        [Paragraph("Approval Status", tbl_label_style), Paragraph(spec_status, tbl_value_style)]
    ]
    meta_tbl = Table(meta_data, colWidths=[180, 324])
    meta_tbl.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, border_color),
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#F0F0F0")),
        ('PADDING', (0,0), (-1,-1), 8),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(meta_tbl)
    story.append(PageBreak())

    # --- Section 1: Executive Summary ---
    story.append(Paragraph("1. Executive Summary", section_heading))
    story.append(Spacer(1, 10))
    story.append(Paragraph(exec_summary or "This document outlines the functional specifications for the approved performance metrics of the organization. Each metric is mapped to strategic objectives and key result areas to ensure operational alignment and governance.", body_style))
    story.append(Spacer(1, 20))

    # --- Section 2: KPI Landscape Overview ---
    story.append(Paragraph("2. KPI Landscape Overview", section_heading))
    story.append(Spacer(1, 10))
    story.append(Paragraph("The following visual mind-map displays the KPI Landscape as a branching tree structure, routing from the core KPI Library through strategic categories down to individual approved metrics.", body_style))
    story.append(Spacer(1, 10))
    
    # Draw KPI Tree
    try:
        story.append(draw_kpi_tree_pdf(items))
        story.append(Spacer(1, 15))
    except Exception as e:
        pass

    story.append(Paragraph("The following table provides a high-level catalog of all approved performance indicators within the scope of this transformation initiative.", body_style))
    story.append(Spacer(1, 10))
    
    land_data = [[
        Paragraph("KPI ID", tbl_hdr_style),
        Paragraph("KPI Name", tbl_hdr_style),
        Paragraph("Category", tbl_hdr_style),
        Paragraph("Functional Area", tbl_hdr_style)
    ]]
    for idx, item in enumerate(items, start=1):
        land_data.append([
            Paragraph(f"KPI-{idx:03d}", tbl_value_style),
            Paragraph(f"<b>{item.kpi_name}</b>", tbl_value_style),
            Paragraph(clean_or_fallback(item.kpi_category, "Operational"), tbl_value_style),
            Paragraph(clean_or_fallback(item.functional_area, "Operations"), tbl_value_style)
        ])
    land_tbl = Table(land_data, colWidths=[80, 184, 120, 120])
    land_tbl.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, border_color),
        ('BACKGROUND', (0,0), (-1,0), dark_gray),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(land_tbl)
    story.append(Spacer(1, 20))

    # --- Section 3: Strategic Traceability Matrix ---
    story.append(Paragraph("3. Strategic Traceability Matrix", section_heading))
    story.append(Spacer(1, 10))
    story.append(Paragraph("This matrix illustrates the strategic alignment from executive objectives down to specific key performance indicators, providing visibility into strategic translation.", body_style))
    story.append(Spacer(1, 10))
    
    trace_data = [[
        Paragraph("Strategic Objective", tbl_hdr_style),
        Paragraph("Business Challenge", tbl_hdr_style),
        Paragraph("KRA", tbl_hdr_style),
        Paragraph("Functional Area", tbl_hdr_style),
        Paragraph("KPI Name", tbl_hdr_style)
    ]]
    for item in items:
        trace_data.append([
            Paragraph(clean_or_fallback(item.strategic_objective_supported, "Optimize Strategy"), tbl_value_style),
            Paragraph(clean_or_fallback(item.business_challenge_addressed, "Inefficient Processes"), tbl_value_style),
            Paragraph(clean_or_fallback(item.related_kra, "Operational Excellence"), tbl_value_style),
            Paragraph(clean_or_fallback(item.functional_area, "Operations"), tbl_value_style),
            Paragraph(f"<b>{item.kpi_name}</b>", tbl_value_style)
        ])
    trace_tbl = Table(trace_data, colWidths=[100, 100, 100, 100, 104])
    trace_tbl.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, border_color),
        ('BACKGROUND', (0,0), (-1,0), dark_gray),
        ('PADDING', (0,0), (-1,-1), 5),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(trace_tbl)
    story.append(PageBreak())

    # --- Section 4: Individual KPI Sections ---
    story.append(Paragraph("4. Individual KPI Specifications", section_heading))
    story.append(Spacer(1, 10))
    
    for idx, item in enumerate(items, start=1):
        kpi_elements = []
        kpi_header_text = f"<font color='#B49600'>Metric {idx:02d}:</font> {item.kpi_name}"
        kpi_elements.append(Paragraph(kpi_header_text, kpi_title_style))
        kpi_elements.append(Spacer(1, 4))
        
        # Meta table for Category, Functional Area, Owners
        kpi_meta_data = [
            [
                Paragraph("KPI Category", tbl_label_style), Paragraph(clean_or_fallback(item.kpi_category, "Operational"), tbl_value_style),
                Paragraph("Functional Area", tbl_label_style), Paragraph(clean_or_fallback(item.functional_area, "Operations"), tbl_value_style)
            ],
            [
                Paragraph("Business Owner", tbl_label_style), Paragraph(clean_or_fallback(item.business_owner, "Advisory Business Owner"), tbl_value_style),
                Paragraph("Data Owner", tbl_label_style), Paragraph(clean_or_fallback(item.data_owner, "Advisory Data Owner"), tbl_value_style)
            ]
        ]
        kpi_meta_tbl = Table(kpi_meta_data, colWidths=[100, 152, 100, 152])
        kpi_meta_tbl.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, border_color),
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#F0F0F0")),
            ('BACKGROUND', (2,0), (2,-1), colors.HexColor("#F0F0F0")),
            ('PADDING', (0,0), (-1,-1), 5),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        kpi_elements.append(kpi_meta_tbl)
        kpi_elements.append(Spacer(1, 8))
        
        # Details: Business Definition & Purpose
        kpi_elements.append(Paragraph("Business Definition", sub_section_title_style))
        kpi_elements.append(Paragraph(clean_or_fallback(item.kpi_definition, "Official definition of the KPI."), body_style))
        
        kpi_elements.append(Paragraph("Business Purpose & Strategic Relevance", sub_section_title_style))
        kpi_elements.append(Paragraph(clean_or_fallback(item.business_purpose_relevance, "Detailing the strategic necessity of tracking this metric."), body_style))
        
        # Calculation Table
        kpi_elements.append(Paragraph("Calculation Methodology", sub_section_title_style))
        calc_rows = [
            [Paragraph("Formula", tbl_label_style), Paragraph(clean_or_fallback(item.formula, "Formula not defined."), tbl_value_style)],
            [Paragraph("Numerator Details", tbl_label_style), Paragraph(clean_or_fallback(item.numerator, "Numerator details."), tbl_value_style)],
            [Paragraph("Denominator Details", tbl_label_style), Paragraph(clean_or_fallback(item.denominator, "Denominator details."), tbl_value_style)],
            [Paragraph("Calculation Logic", tbl_label_style), Paragraph(clean_or_fallback(item.calculation_methodology, "Standard calculation procedures."), tbl_value_style)],
            [Paragraph("Inclusion Rules", tbl_label_style), Paragraph(clean_or_fallback(item.inclusion_rules, "Include all standard transactions."), tbl_value_style)],
            [Paragraph("Exclusion Rules", tbl_label_style), Paragraph(clean_or_fallback(item.exclusion_rules, "Exclude test and non-operational entries."), tbl_value_style)],
            [Paragraph("Sample Worked Example", tbl_label_style), Paragraph(clean_or_fallback(item.sample_calculation, "Sample: (Subset / Total) * 100"), tbl_value_style)]
        ]
        calc_tbl = Table(calc_rows, colWidths=[130, 374])
        calc_tbl.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, border_color),
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#F9F9F9")),
            ('PADDING', (0,0), (-1,-1), 5),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ]))
        kpi_elements.append(calc_tbl)
        kpi_elements.append(Spacer(1, 8))
        
        # Business Rules & Validation Table
        kpi_elements.append(Paragraph("Business Rules & Data Validation", sub_section_title_style))
        rules_rows = [
            [Paragraph("Business Operating Rules", tbl_label_style), Paragraph(clean_or_fallback(item.business_rules, "Standard business operating rules."), tbl_value_style)],
            [Paragraph("Data Validation Rules", tbl_label_style), Paragraph(clean_or_fallback(item.data_validation_rules, "Ensure numeric bounds and data integrity."), tbl_value_style)],
            [Paragraph("Exception Handling Rules", tbl_label_style), Paragraph(clean_or_fallback(item.exception_handling_rules, "Standard null/zero denominator exceptions."), tbl_value_style)],
            [Paragraph("Data Quality Expectations", tbl_label_style), Paragraph(clean_or_fallback(item.data_quality_expectations, "High accuracy in source ledger."), tbl_value_style)]
        ]
        rules_tbl = Table(rules_rows, colWidths=[130, 374])
        rules_tbl.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, border_color),
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#F9F9F9")),
            ('PADDING', (0,0), (-1,-1), 5),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ]))
        kpi_elements.append(rules_tbl)
        kpi_elements.append(Spacer(1, 8))
        
        # Source Systems
        kpi_elements.append(Paragraph("Recommended Source Systems", sub_section_title_style))
        kpi_elements.append(Paragraph(f"<i>{clean_or_fallback(item.source_systems_lineage, 'Recommended source ERP database. (Assumption)')}</i>", body_style))
        kpi_elements.append(Spacer(1, 20))
        
        story.append(KeepTogether(kpi_elements))
    
    story.append(PageBreak())

    # --- Section 5: Governance Framework ---
    story.append(Paragraph("5. Governance Framework", section_heading))
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "To ensure metric consistency, accountability, and ongoing relevance, a formal governance framework is "
        "established for all approved KPIs. This structure assigns clear responsibilities and defines escalation paths.",
        body_style
    ))
    story.append(Spacer(1, 8))
    
    story.append(Paragraph("Roles and Responsibilities", sub_section_title_style))
    story.append(Paragraph("&bull; <b>Business Owner:</b> Responsible for defining the business logic, validating calculation results, approving target thresholds, and driving operational performance based on metric insights.", bullet_style))
    story.append(Paragraph("&bull; <b>Data Owner:</b> Accountable for technical lineage, data completeness, source-to-target mapping, ETL data quality checks, and resolving data ingestion or availability issues.", bullet_style))
    story.append(Paragraph("&bull; <b>Escalation Path:</b> In case of data quality discrepancies or alignment disputes, issues are escalated to the Data Governance Committee and KPI Advisory Board for review and reconciliation.", bullet_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("Ownership and Governance Policy Matrix", sub_section_title_style))
    story.append(Spacer(1, 6))
    gov_rows = [[
        Paragraph("Metric Name", tbl_hdr_style),
        Paragraph("Business Owner", tbl_hdr_style),
        Paragraph("Data Owner", tbl_hdr_style),
        Paragraph("Governance Policy Notes", tbl_hdr_style)
    ]]
    for item in items:
        gov_rows.append([
            Paragraph(f"<b>{item.kpi_name}</b>", tbl_value_style),
            Paragraph(clean_or_fallback(item.business_owner, "Business Sponsor"), tbl_value_style),
            Paragraph(clean_or_fallback(item.data_owner, "Data Custodian"), tbl_value_style),
            Paragraph(clean_or_fallback(item.ownership_governance, "Subject to standard quarterly advisory audits and performance reviews."), tbl_value_style)
        ])
    gov_tbl = Table(gov_rows, colWidths=[110, 100, 100, 194])
    gov_tbl.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, border_color),
        ('BACKGROUND', (0,0), (-1,0), dark_gray),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(gov_tbl)
    story.append(PageBreak())

    # --- Section 6: Reporting & Dashboard Requirements ---
    story.append(Paragraph("6. Reporting & Dashboard Requirements", section_heading))
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "Visualization layout and threshold criteria dictate how data is displayed to support decision-making. "
        "The matrix below outlines reporting recommendations and performance thresholds for each metric.",
        body_style
    ))
    story.append(Spacer(1, 10))
    
    rep_rows = [[
        Paragraph("Metric Name", tbl_hdr_style),
        Paragraph("Reporting Guidelines", tbl_hdr_style),
        Paragraph("Dashboard Placement", tbl_hdr_style),
        Paragraph("Threshold Guidance", tbl_hdr_style)
    ]]
    for item in items:
        def_rep = get_default_reporting(item.kpi_category)
        def_thresh = get_default_threshold(item.kpi_category)
        def_dash = "Standard Performance Dashboard."
        
        rep_guidelines = clean_or_fallback(item.reporting_requirements, def_rep)
        dash_placement = clean_or_fallback(item.dashboard_recommendations, def_dash)
        threshold_guidance = clean_or_fallback(item.threshold_guidance, def_thresh)
        
        rep_rows.append([
            Paragraph(f"<b>{item.kpi_name}</b>", tbl_value_style),
            Paragraph(rep_guidelines, tbl_value_style),
            Paragraph(dash_placement, tbl_value_style),
            Paragraph(threshold_guidance, tbl_value_style)
        ])
    rep_tbl = Table(rep_rows, colWidths=[100, 140, 110, 154])
    rep_tbl.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, border_color),
        ('BACKGROUND', (0,0), (-1,0), dark_gray),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(rep_tbl)
    story.append(PageBreak())

    # --- Section 7: Assumptions & Constraints ---
    story.append(Paragraph("7. Assumptions & Constraints", section_heading))
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "A clear understanding of business and technical assumptions is critical for successful implementation. "
        "The following list represents consolidated baseline assumptions and constraints.",
        body_style
    ))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("General Architectural Assumptions", sub_section_title_style))
    story.append(Paragraph("&bull; <b>Data Availability:</b> Source transactional tables are assumed to be loaded into the central reporting repository on a standard nightly batch cadence.", bullet_style))
    story.append(Paragraph("&bull; <b>Fiscal Calendar:</b> Standard calendar year rules are assumed unless otherwise explicitly documented in specific financial indicators.", bullet_style))
    story.append(Paragraph("&bull; <b>System Uptime:</b> Target source ERP ledgers are assumed to maintain 99.5% uptime during standard reporting extraction windows.", bullet_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("Metric-Specific Assumptions and Limitations", sub_section_title_style))
    for item in items:
        cleaned_asm = clean_or_fallback(item.assumptions_constraints, "").strip()
        if cleaned_asm:
            story.append(Paragraph(f"&bull; <b>{item.kpi_name}:</b> {cleaned_asm}", bullet_style))
            
    story.append(PageBreak())

    # --- Section 8: Implementation Considerations ---
    story.append(Paragraph("8. Implementation Considerations", section_heading))
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "Transitioning these specifications into functional BI tools requires rigorous testing, change management, "
        "and data reconciliation. Standard implementation considerations are outlined below.",
        body_style
    ))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("Standard Implementation Guidelines", sub_section_title_style))
    story.append(Paragraph("&bull; <b>Data Reconciliation:</b> All KPI calculation outcomes must be audited and reconciled against official books of record or audited financial statements.", bullet_style))
    story.append(Paragraph("&bull; <b>User Acceptance Testing (UAT):</b> Business owners must perform visual and numeric validation of dashboard mockups prior to production sign-off.", bullet_style))
    story.append(Paragraph("&bull; <b>Change Management:</b> Training sessions and clear system documentation are required to support user onboarding and ensure high organizational adoption.", bullet_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("Metric-Specific Technical Considerations", sub_section_title_style))
    for item in items:
        cleaned_imp = clean_or_fallback(item.implementation_guidance, "").strip()
        if cleaned_imp:
            story.append(Paragraph(f"&bull; <b>{item.kpi_name}:</b> {cleaned_imp}", bullet_style))
            
    story.append(PageBreak())

    # --- Section 9: Appendix ---
    story.append(Paragraph("9. Appendix", section_heading))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("KPI Glossary", sub_section_title_style))
    glossary_items = [
        ("KPI (Key Performance Indicator)", "A quantifiable measure used to evaluate the success of an organization or activity in meeting performance objectives."),
        ("KRA (Key Result Area)", "Primary focus areas of outcomes or outputs for which an organizational unit or role is responsible."),
        ("Numerator", "The upper portion of a division representing the measured subset of occurrences."),
        ("Denominator", "The lower portion of a division representing the total base population.")
    ]
    for term, defn in glossary_items:
        story.append(Paragraph(f"&bull; <b>{term}:</b> {defn}", bullet_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("Data Quality Principles", sub_section_title_style))
    dq_items = [
        ("Accuracy", "Data correctly represents the real-world operational event it records."),
        ("Completeness", "All necessary dataset components are present without omission."),
        ("Consistency", "Metrics align across various systems, business units, and report interfaces."),
        ("Timeliness", "Updates occur within the required reporting cadence and operational windows.")
    ]
    for principle, desc in dq_items:
        story.append(Paragraph(f"&bull; <b>{principle}:</b> {desc}", bullet_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("Acronym Reference", sub_section_title_style))
    ac_items = [
        ("ERP", "Enterprise Resource Planning"),
        ("SAP FI-CO", "Financial Accounting & Controlling"),
        ("SAP SD", "Sales & Distribution"),
        ("SAP MM", "Materials Management"),
        ("BI / DWH", "Business Intelligence / Data Warehouse"),
        ("UAT", "User Acceptance Testing")
    ]
    for acr, full in ac_items:
        story.append(Paragraph(f"&bull; <b>{acr}:</b> {full}", bullet_style))
        
    doc.build(story, canvasmaker=NumberedCanvas)



def generate_docx_prompt(path: Path, prompt_text: str, context: BusinessContext) -> None:
    """Generates a professional Word Document for the KPI Prompt."""
    doc = Document()
    
    # Configure page margins (1.0 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(30, 30, 30)

    # Title Page / Header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("EY KPI Advisory & Analytics Copilot")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(180, 150, 0) # EY Gold

    main_title_p = doc.add_paragraph()
    main_title_p.paragraph_format.space_after = Pt(12)
    main_title_run = main_title_p.add_run("KPI Generation Prompt & Requirements")
    main_title_run.font.name = 'Calibri'
    main_title_run.font.size = Pt(22)
    main_title_run.font.bold = True
    main_title_run.font.color.rgb = RGBColor(27, 27, 27)

    # Metadata Block
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(24)
    meta_p.paragraph_format.line_spacing = 1.3
    meta_p.add_run("Industry Sector: ").bold = True
    meta_p.add_run(f"{context.industry}\n")
    meta_p.add_run("Organizational Level: ").bold = True
    meta_p.add_run(f"{context.organization_level}\n")
    meta_p.add_run("Strategic KPI Count: ").bold = True
    meta_p.add_run(f"{context.kpi_count} Metrics\n")

    # Yellow accent separator
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.autofit = False
    accent_table.columns[0].width = Inches(6.5)
    accent_cell = accent_table.cell(0, 0)
    set_cell_shading(accent_cell, "FFE600") # EY Yellow
    set_cell_margins(accent_cell, top=20, bottom=20, left=0, right=0)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Split prompt_text into lines and parse
    in_code_block = False
    lines = prompt_text.splitlines()
    for line in lines:
        cleaned_line = line.strip()
        
        # Code block toggle
        if cleaned_line.startswith("```"):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
            continue

        if not cleaned_line:
            # Add small spacing for empty lines, but avoid adding empty paragraphs
            continue

        # Markdown headings
        if line.startswith("# "):
            heading_text = line[2:].strip()
            heading = doc.add_heading(level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(4)
            heading.paragraph_format.keep_with_next = True
            run = heading.add_run(heading_text)
            run.font.name = 'Calibri'
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 27, 27)
        elif line.startswith("## "):
            heading_text = line[3:].strip()
            heading = doc.add_heading(level=2)
            heading.paragraph_format.space_before = Pt(10)
            heading.paragraph_format.space_after = Pt(3)
            heading.paragraph_format.keep_with_next = True
            run = heading.add_run(heading_text)
            run.font.name = 'Calibri'
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 27, 27)
        elif line.startswith("### "):
            heading_text = line[4:].strip()
            heading = doc.add_heading(level=3)
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after = Pt(2)
            heading.paragraph_format.keep_with_next = True
            run = heading.add_run(heading_text)
            run.font.name = 'Calibri'
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 27, 27)
        # Bullet points
        elif line.startswith("- ") or line.startswith("* "):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(bullet_text)
            run.font.name = 'Calibri'
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(line)
            run.font.name = 'Calibri'

    doc.save(path)

